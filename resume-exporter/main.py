#!/usr/bin/env python3
# resume-exporter.py

import os
import json
import pika
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# adjust path to your utils directory
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../utils')))
from utils import get_resumes_collection

# ─── Configuration ─────────────────────────────────────────────────────────────
MONGO_COLLECTION = get_resumes_collection()

RABBIT_CONFIG = {
    'host': 'localhost',
    'queues': {
        'tailored': 'resume.tailored',
        'completed': 'resume.completed'
    }
}

OUTPUT_DIR = 'docs'
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── RabbitMQ Setup ─────────────────────────────────────────────────────────────
connection = pika.BlockingConnection(pika.ConnectionParameters(RABBIT_CONFIG['host']))
channel = connection.channel()
channel.queue_declare(queue=RABBIT_CONFIG['queues']['tailored'])
channel.queue_declare(queue=RABBIT_CONFIG['queues']['completed'])

def generate_docx(data: dict, output_path: str) -> None:
    """
    Renders all sections of the resume JSON into a DOCX with Calibri font,
    section titles at 11pt, body text at 10.5pt, no spacing between paragraphs,
    and narrow page margins.
    """
    doc = Document()

    # ── Page Margins (narrow) ─────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ── Default Style: Calibri 10.5pt, zero spacing ──────────────────────
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)

    # Helper to reset spacing on a paragraph
    def no_space(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    # ── Header: Name (11pt) ─────────────────────────────────────────────
    name_h = doc.add_heading(level=0)
    run = name_h.add_run(data.get('name', ''))
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    no_space(name_h)

    # ── Contact Info (centered, 10.5pt) ─────────────────────────────────
    contact_p = doc.add_paragraph()
    no_space(contact_p)
    contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_text = f"{data.get('contact_number','')} | {data.get('email','')} | {data.get('linkedin','')}"
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(10.5)

    # ── Section Title Helper ────────────────────────────────────────────
    def add_section(title: str):
        p = doc.add_paragraph()
        no_space(p)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    # ── Skills ──────────────────────────────────────────────────────────
    add_section('Skills')
    for cat, skills in data.get('skills', {}).items():
        p = doc.add_paragraph(style='List Bullet')
        no_space(p)
        run = p.add_run(f"{cat.replace('_',' ')}: ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        rest = p.add_run(skills)
        rest.font.name = 'Calibri'
        rest.font.size = Pt(10.5)

    # ── Experience ───────────────────────────────────────────────────────
    add_section('Experience')
    for exp in data.get('experience', []):
        # Company + Date
        p = doc.add_paragraph()
        no_space(p)
        run = p.add_run(exp.get('company',''))
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        date_run = p.add_run(f"\t{exp.get('date','')}")
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(10.5)

        # Role
        role_p = doc.add_paragraph()
        no_space(role_p)
        role_run = role_p.add_run(exp.get('role',''))
        role_run.italic = True
        role_run.font.name = 'Calibri'
        role_run.font.size = Pt(10.5)

        # Responsibilities
        for b in exp.get('responsibilities', []):
            bullet = doc.add_paragraph(style='List Bullet')
            no_space(bullet)
            bullet.level = 1
            brun = bullet.add_run(b)
            brun.font.name = 'Calibri'
            brun.font.size = Pt(10.5)

    # ── Education ────────────────────────────────────────────────────────
    add_section('Education')
    for edu in data.get('education', []):
        p = doc.add_paragraph()
        no_space(p)
        text = f"{edu.get('institution','')} | {edu.get('degree','')}, GPA: {edu.get('gpa','')}"
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        date_run = p.add_run(f"\t{edu.get('date','')}")
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(10.5)

    # ── Projects ─────────────────────────────────────────────────────────
    add_section('Projects')
    for proj in data.get('projects', []):
        p = doc.add_paragraph()
        no_space(p)
        run = p.add_run(f"{proj.get('title','')} – {proj.get('date','')}")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        for d in proj.get('details', []):
            bullet = doc.add_paragraph(style='List Bullet')
            no_space(bullet)
            bullet.level = 1
            brun = bullet.add_run(d)
            brun.font.name = 'Calibri'
            brun.font.size = Pt(10.5)

    # ── Honors & Awards ──────────────────────────────────────────────────
    add_section('Honors and Awards')
    for award in data.get('honors_and_awards', []):
        p = doc.add_paragraph(style='List Bullet')
        no_space(p)
        run = p.add_run(award)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)

def on_tailored(ch, method, properties, body):
    msg = json.loads(body)
    resume_id = msg.get('resume_id')
    print(f"[Exporter] Received tailored resume event: {resume_id}")

    doc = MONGO_COLLECTION.find_one({'_id': resume_id})
    if not doc or 'tailored_resume' not in doc:
        print(f"[Exporter] ERROR: no tailored_resume for {resume_id}")
        return

    resume_json = json.loads(doc['tailored_resume'])

    # **BUILD A FULL FILEPATH**, not just the directory!
    output_path = os.path.join(OUTPUT_DIR, f"{resume_id}.docx")
    generate_docx(resume_json, output_path)
    print(f"[Exporter] DOCX generated at: {output_path}")

   

# Start consuming tailored events
channel.basic_consume(
    queue=RABBIT_CONFIG['queues']['tailored'],
    on_message_callback=on_tailored,
    auto_ack=True
)

print("Resume Exporter Service: awaiting messages on", RABBIT_CONFIG['queues']['tailored'])
channel.start_consuming()
